"""
Document Exporter Module

This module provides functionality to export content to various document
formats, including DOCX, PDF, and TXT. It handles document creation,
formatting, and export operations.
"""

from pathlib import Path
from typing import Any, Dict, List, Optional, Union, Literal

# Optional imports with type ignore for mypy
# These will be checked at runtime

try:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
    )  # type: ignore
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class DocExporter:
    """A class to handle document export to various formats.

    This class provides methods to export content to different document formats
    including DOCX, PDF, and TXT. It supports rich text, images, and basic formatting.
    """

    def __init__(self, output_dir: Optional[Union[str, Path]] = None) -> None:
        """Initialize the DocExporter with an optional output directory.

        Args:
            output_dir: Directory to save exported files
                      (default: current directory)
        """
        self.output_dir = Path(output_dir) if output_dir else Path.cwd()
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Check for required dependencies and raise informative errors."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install with: pip install python-docx"
            )
        if not PDF_AVAILABLE:
            print(
                "Warning: reportlab is not installed. "
                "PDF export will not be available. "
                "Install with: pip install reportlab"
            )

    def export_docx(
        self,
        content: Union[str, List[Dict[str, Any]]],
        output_file: Optional[Union[str, Path]] = None,
        title: Optional[str] = None,
        author: str = "SVGG Exporter",
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Export content to a DOCX file.

        Args:
            content: Content to export (can be plain text, HTML,
                    or structured data)
            output_file: Output file path (default: based on title or timestamp)
            title: Document title
            author: Document author
            **kwargs: Additional formatting options

        Returns:
            Dictionary with export results and metadata
        """
        if not DOCX_AVAILABLE:
            return {
                'success': False,
                'error': 'python-docx is not installed',
                'suggestion': 'Install with: pip install python-docx',
            }

        try:
            doc = Document()

            # Set document properties
            core_props = doc.core_properties
            core_props.author = author
            if title:
                core_props.title = title
                doc.add_heading(title, level=0)
                doc.add_paragraph()  # Add some space after title

            # Handle different content types
            if isinstance(content, str):
                # Simple text content
                doc.add_paragraph(content)
            elif isinstance(content, list):
                # Structured content with formatting
                for item in content:
                    self._add_docx_element(doc, item)

            # Save the document
            if not output_file:
                filename = f"document_{title or 'export'}.docx"
                filename = "".join(c if c.isalnum() else "_" for c in filename)
                output_file = self.output_dir / filename
            else:
                output_file = Path(output_file)

            doc.save(str(output_file))

            return {
                'success': True,
                'output_file': str(output_file),
                'size': output_file.stat().st_size,
                'format': 'docx',
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
            }

    def _add_docx_element(
        self,
        doc: Any,
        element: Dict[str, Any],
    ) -> None:
        """Add a structured element to a Word document.

        Args:
            doc: python-docx Document object
            element: Dictionary describing the element to add
        """
        element_type = element.get('type', 'paragraph')
        text = element.get('text', '')
        style = element.get('style', 'Normal')
        level = element.get('level', 1)

        if element_type == 'heading':
            doc.add_heading(text, level=level)
        elif element_type == 'paragraph':
            p = doc.add_paragraph(style=style)
            p.add_run(text)
        elif element_type == 'list':
            items = element.get('items', [])
            for item in items:
                p = doc.add_paragraph(style='ListBullet')
                p.add_run(str(item))
        elif element_type == 'image' and 'path' in element:
            try:
                doc.add_picture(
                    element['path'],
                    width=Inches(element.get('width', 6)),
                )
            except Exception:  # noqa: BLE001
                # Skip if image can't be added
                pass

    def export_pdf(
        self,
        content: Union[str, List[Dict[str, Any]]],
        output_file: Optional[Union[str, Path]] = None,
        title: Optional[str] = None,
        author: str = "SVGG Exporter",
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Export content to a PDF file.

        Args:
            content: Content to export (can be plain text, HTML,
                    or structured data)
            output_file: Output file path (default: based on title or timestamp)
            title: Document title
            author: Document author
            **kwargs: Additional formatting options

        Returns:
            Dictionary with export results and metadata
        """
        if not PDF_AVAILABLE:
            return {
                'success': False,
                'error': 'reportlab is not installed',
                'suggestion': 'Install with: pip install reportlab',
            }

        try:
            # Generate output filename if not provided
            if not output_file:
                filename = f"document_{title or 'export'}.pdf"
                filename = "".join(c if c.isalnum() else "_" for c in filename)
                output_file = self.output_dir / filename
            else:
                output_file = Path(output_file)

            # Create PDF document
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72,
            )

            # Prepare content
            styles = getSampleStyleSheet()
            story = []

            # Add title if provided
            if title:
                title_style = styles['Heading1']
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 12))

            # Add content
            if isinstance(content, str):
                # Simple text content
                story.append(Paragraph(content, styles['Normal']))
            elif isinstance(content, list):
                # Structured content
                for item in content:
                    self._add_pdf_element(story, item, styles)

            # Build the PDF
            doc.build(story)

            return {
                'success': True,
                'output_file': str(output_file),
                'size': output_file.stat().st_size,
                'format': 'pdf',
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
            }

    def _add_pdf_element(
        self,
        story: List[Any],
        element: Dict[str, Any],
        styles: Any,
    ) -> None:
        """Add a structured element to a PDF document.

        Args:
            story: List of elements in the PDF story
            element: Dictionary describing the element to add
            styles: ReportLab styles
        """
        element_type = element.get('type', 'paragraph')
        text = element.get('text', '')

        if element_type == 'heading':
            level = element.get('level', 1)
            style_name = f'Heading{min(level, 6)}'
            story.append(Paragraph(text, styles[style_name]))
        elif element_type == 'paragraph':
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 12))
        elif element_type == 'list':
            items = element.get('items', [])
            for item in items:
                story.append(Paragraph(f"• {item}", styles['Normal']))
                story.append(Spacer(1, 6))

    def export_txt(
        self,
        content: Union[str, List[Dict[str, Any]]],
        output_file: Optional[Union[str, Path]] = None,
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Export content to a plain text file.

        Args:
            content: Content to export (can be plain text or structured data)
            output_file: Output file path (default: based on timestamp)
            **kwargs: Additional options (ignored for TXT export)

        Returns:
            Dictionary with export results and metadata
        """
        try:
            # Generate output filename if not provided
            if not output_file:
                filename = "document_export.txt"
                output_file = self.output_dir / filename
            else:
                output_file = Path(output_file)

            # Convert content to text
            text_content = self._convert_to_text(content)

            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text_content)

            return {
                'success': True,
                'output_file': str(output_file),
                'size': output_file.stat().st_size,
                'format': 'txt',
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
            }

    def _convert_to_text(self, content: Union[str, List[Dict[str, Any]]]) -> str:
        """Convert content to plain text.

        Args:
            content: Content to convert (string or structured data)

        Returns:
            Plain text representation of the content
        """
        if isinstance(content, str):
            return content

        if not isinstance(content, list):
            return str(content)

        lines = []
        for item in content:
            element_type = item.get('type', 'text')
            text = item.get('text', '')

            if element_type == 'heading':
                equals = '=' * len(text)
                heading = f'\n{text.upper()}\n{equals}\n'
                lines.append(heading)
            elif element_type == 'paragraph':
                lines.append(f'{text}\n')
            elif element_type == 'list':
                items = item.get('items', [])
                for i, item_text in enumerate(items, 1):
                    lines.append(f"{i}. {item_text}")
                lines.append('')  # Add empty line after list
            elif element_type == 'image' and 'path' in item:
                lines.append(f"[Image: {item['path']}]\n")

        return '\n'.join(lines).strip()

    def export(
        self,
        content: Union[str, List[Dict[str, Any]]],
        format: Literal['docx', 'pdf', 'txt'] = 'docx',
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Export content to the specified format.

        Args:
            content: Content to export
            format: Output format (docx, pdf, or txt)
            **kwargs: Additional options passed to the specific exporter

        Returns:
            Dictionary with export results and metadata
        """
        fmt = format.lower()
        if fmt == 'docx':
            return self.export_docx(content, **kwargs)
        if fmt == 'pdf':
            return self.export_pdf(content, **kwargs)
        if fmt == 'txt':
            return self.export_txt(content, **kwargs)
        return {
            'success': False,
            'error': f'Unsupported format: {format}',
            'supported_formats': ['docx', 'pdf', 'txt'],
        }

    def __enter__(self) -> 'DocExporter':
        """Context manager entry."""
        return self

    def __exit__(self, exc_type: Any, exc_val: Any, exc_tb: Any) -> None:
        """Context manager exit - clean up resources."""
        pass  # No resources to clean up
